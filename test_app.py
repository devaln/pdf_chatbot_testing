# --- Imports ---
import os, json, uuid, shutil, re, sys
from pathlib import Path
import concurrent.futures
import pytesseract
import pandas as pd
from PIL import Image
from pdf2image import convert_from_path
import numpy as np
from difflib import SequenceMatcher
import streamlit as st
from paddleocr import PaddleOCR
from docx import Document as DocxDocument
import docx
import io

from langchain.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.chat_models import ChatOllama
from langchain_core.prompts import ChatPromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain_experimental.agents import create_pandas_dataframe_agent
from langchain.agents.agent_types import AgentType

# --- Config ---
OLLAMA_BASE_URL = "http://localhost:11434"
OLLAMA_LLM_MODEL = "llama3:latest"
HF_EMBED_MODEL = "intfloat/e5-base"
DB_DIR = "./faiss_db"
TABLE_DIR = "./tables"
TEMP_DIR = "./temp"
TOP_K = 5
TEXT_CHUNK_SIZE = 500

# --- Setup ---
os.makedirs(TABLE_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)
embedder = HuggingFaceEmbeddings(model_name=HF_EMBED_MODEL)
ocr = PaddleOCR(use_angle_cls=True, lang='en')

if "msgs" not in st.session_state: st.session_state.msgs = []
if "vs" not in st.session_state: st.session_state.vs = None
if "uploaded_files" not in st.session_state: st.session_state.uploaded_files = []
if "chat_history" not in st.session_state: st.session_state.chat_history = []
if "memory" not in st.session_state:
    st.session_state.memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

# --- Extractors ---
def extract_from_pdf(pdf_path):
    docs = []
    fname = os.path.basename(pdf_path)
    images = convert_from_path(pdf_path, dpi=300)
    for pn, img in enumerate(images):
        img = img.convert("RGB")
        result = ocr.ocr(np.array(img), cls=True)
        text_lines = [line[1][0].strip() for line in result[0] if line[1][0].strip()]
        full_text = "\n".join(text_lines)
        for i in range(0, len(full_text), TEXT_CHUNK_SIZE):
            chunk = full_text[i:i+TEXT_CHUNK_SIZE].strip()
            if chunk:
                docs.append(Document(page_content=chunk, metadata={"source": fname, "page": pn+1, "type": "text"}))
    return docs

def extract_from_docx(file_bytes, fname):
    docs = []
    doc = DocxDocument(io.BytesIO(file_bytes))
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    for i in range(0, len(full_text), TEXT_CHUNK_SIZE):
        chunk = full_text[i:i+TEXT_CHUNK_SIZE].strip()
        if chunk:
            docs.append(Document(page_content=chunk, metadata={"source": fname, "type": "docx"}))
    return docs

def extract_from_csv(file_bytes, fname):
    docs = []
    df = pd.read_csv(io.BytesIO(file_bytes))
    rows = df.astype(str).values.tolist()
    full_text = "\n".join([", ".join(row) for row in rows])
    for i in range(0, len(full_text), TEXT_CHUNK_SIZE):
        chunk = full_text[i:i+TEXT_CHUNK_SIZE].strip()
        if chunk:
            docs.append(Document(page_content=chunk, metadata={"source": fname, "type": "csv"}))
    df.to_json(f"{TABLE_DIR}/{uuid.uuid4()}.json", orient="records")
    return docs

def extract_from_excel(file_bytes, fname):
    docs = []
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        rows = df.astype(str).values.tolist()
        full_text = "\n".join([", ".join(row) for row in rows])
        for i in range(0, len(full_text), TEXT_CHUNK_SIZE):
            chunk = full_text[i:i+TEXT_CHUNK_SIZE].strip()
            if chunk:
                docs.append(Document(page_content=chunk, metadata={"source": fname, "sheet": sheet, "type": "excel"}))
        df.to_json(f"{TABLE_DIR}/{uuid.uuid4()}.json", orient="records")
    return docs

# --- Indexing ---
def process_and_index(files):
    docs = []
    for file in files:
        fname = file.name
        fpath = os.path.join(TEMP_DIR, fname)
        with open(fpath, "wb") as f:
            f.write(file.getbuffer())

        ext = fname.lower().split(".")[-1]
        fbytes = file.getvalue()
        if ext == "pdf":
            docs.extend(extract_from_pdf(fpath))
        elif ext == "docx":
            docs.extend(extract_from_docx(fbytes, fname))
        elif ext == "csv":
            docs.extend(extract_from_csv(fbytes, fname))
        elif ext in ["xls", "xlsx"]:
            docs.extend(extract_from_excel(fbytes, fname))

    if not docs: return None
    if Path(DB_DIR).exists():
        vs = FAISS.load_local(DB_DIR, embedder, allow_dangerous_deserialization=True)
        vs.add_documents(docs)
    else:
        vs = FAISS.from_documents(docs, embedder)
    vs.save_local(DB_DIR)
    shutil.rmtree(TEMP_DIR, ignore_errors=True)
    return vs

# --- LangChain Chain ---
def get_conversational_chain(vs):
    retriever = vs.as_retriever(search_kwargs={"k": TOP_K})
    prompt = ChatPromptTemplate.from_template("""
You are a helpful assistant. Cite document name and page/sheet. Use prior chat history to understand follow-ups.

Chat History:
{chat_history}

Context from documents:
{context}

Question: {question}
Answer:
""")
    llm = ChatOllama(model=OLLAMA_LLM_MODEL, base_url=OLLAMA_BASE_URL, temperature=0.1)
    memory = st.session_state.memory
    return ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=retriever,
        memory=memory,
        combine_docs_chain_kwargs={"prompt": prompt}
    )

def smart_split(query):
    return [q.strip() for q in re.split(r'\s*&\s*', query) if q.strip()]

def route_query(q, vs):
    if any(w in q.lower() for w in ["sum", "average", "mean", "trend", "total", "compare"]):
        tables = sorted(Path(TABLE_DIR).glob("*.json"), key=os.path.getmtime, reverse=True)
        if not tables:
            return "\u26a0 No table/chart data available."
        df = pd.read_json(tables[0])
        agent = create_pandas_dataframe_agent(
            ChatOllama(model=OLLAMA_LLM_MODEL, base_url=OLLAMA_BASE_URL),
            df,
            agent_type=AgentType.OPENAI_FUNCTIONS,
            verbose=False
        )
        return agent.run(q)
    else:
        chain = get_conversational_chain(vs)
        return chain.run({"question": q})

# --- Streamlit UI ---
st.set_page_config(page_title="Doc Q&A", layout="wide")
st.title("Multi-Doc Q&A with PaddleOCR")

with st.sidebar:
    st.header("Upload Documents")
    uploader = st.file_uploader("Upload PDF, DOCX, Excel, or CSV", type=["pdf", "docx", "xlsx", "xls", "csv"], accept_multiple_files=True)
    if uploader:
        st.session_state.uploaded_files = uploader
        for f in uploader:
            st.markdown(f"\u2705 {f.name}")

    if st.button("Extract & Index"):
        if st.session_state.uploaded_files:
            with st.spinner("Extracting..."):
                vs = process_and_index(st.session_state.uploaded_files)
                if vs:
                    st.session_state.vs = vs
                    st.session_state.msgs = []
                    st.success("Indexing complete!")
                    st.session_state.uploaded_files = []
        else:
            st.warning("Please upload documents.")

    if st.button("Clear Chat"): st.session_state.msgs = []
    if st.button("Clear FAISS Index"):
        shutil.rmtree(DB_DIR, ignore_errors=True)
        st.session_state.vs = None
        st.success("Index deleted.")
    if st.button("Clear Memory"):
        st.session_state.memory.clear()
        st.session_state.chat_history = []

if st.session_state.vs is None and Path(DB_DIR).exists():
    st.session_state.vs = FAISS.load_local(DB_DIR, embedder, allow_dangerous_deserialization=True)

st.markdown("### Ask your question")
query = st.chat_input("Ask questions (e.g. 'total revenue & trend in sales')")

if query:
    st.session_state.msgs.append({"role": "user", "content": query})
    with st.spinner("Processing..."):
        responses = []
        sub_qs = smart_split(query)
        with concurrent.futures.ThreadPoolExecutor() as ex:
            future_map = {ex.submit(route_query, q, st.session_state.vs): q for q in sub_qs}
            for idx, future in enumerate(concurrent.futures.as_completed(future_map)):
                sub_q = future_map[future]
                try:
                    result = future.result()
                    responses.append(f"### Q{idx+1}: {sub_q}\n\n{result}")

                    qa_doc = Document(
                        page_content=f"Q: {sub_q}\nA: {result}",
                        metadata={"source": "chat_memory", "type": "chat", "user": "default"}
                    )
                    st.session_state.vs.add_documents([qa_doc])
                    st.session_state.vs.save_local(DB_DIR)
                except Exception as e:
                    responses.append(f"### Q{idx+1}: {sub_q}\n\n\u26a0 Error: {e}")
        st.session_state.msgs.append({"role": "assistant", "content": "\n\n---\n\n".join(responses)})

for msg in st.session_state.msgs:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"], unsafe_allow_html=True)